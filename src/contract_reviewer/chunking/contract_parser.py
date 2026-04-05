"""Parse contract documents (.pdf, .docx, .txt) into structured sections."""

import logging
import re
from pathlib import Path

from contract_reviewer.models.contract import Contract, Section

logger = logging.getLogger(__name__)


class ParseError(Exception):
    """Raised when a contract file cannot be parsed."""


class ContractParser:
    """Extract structured sections from contract documents."""

    # Common Chinese contract section patterns
    SECTION_PATTERNS = [
        re.compile(r"^第[一二三四五六七八九十百]+[条章节编]"),  # 第一条, 第二章
        re.compile(r"^[一二三四五六七八九十]+[、.]"),  # 一、 二.
        re.compile(r"^[（(]\s*[一二三四五六七八九十]+\s*[）)]"),  # （一） (二)
        re.compile(r"^\d+[\.、]\s"),  # 1. 2、
        re.compile(r"^Article\s+\d+", re.IGNORECASE),  # Article 1
        re.compile(r"^Section\s+\d+", re.IGNORECASE),  # Section 1
    ]

    # Minimum extractable text length to consider a PDF valid
    MIN_TEXT_LENGTH = 50

    def __init__(self, ocr_engine: "OCREngine | None" = None):
        """Initialize parser with optional OCR engine for scanned PDFs.

        Args:
            ocr_engine: An OCR engine instance. If provided, scanned PDFs
                will be processed via OCR instead of raising an error.
        """
        self.ocr_engine = ocr_engine

    async def parse_async(self, file_path: str) -> Contract:
        """Parse a contract file (async version, required for OCR fallback)."""
        path = Path(file_path)
        if not path.exists():
            raise ParseError(f"File not found: {file_path}")

        text = self._extract_text(path)
        metadata: dict = {}

        # If text extraction failed and OCR is available, try OCR
        if len(text.strip()) < self.MIN_TEXT_LENGTH and path.suffix.lower() == ".pdf":
            if self.ocr_engine:
                logger.info("Text extraction insufficient, attempting OCR: %s", path.name)
                ocr_result = await self.ocr_engine.recognize(str(path))
                text = ocr_result.full_text
                metadata = {
                    "ocr_used": True,
                    "ocr_provider": ocr_result.provider,
                    "ocr_device": ocr_result.device,
                    "ocr_confidence": ocr_result.average_confidence,
                    "ocr_elapsed_ms": ocr_result.total_elapsed_ms,
                }
                logger.info(
                    "OCR complete: provider=%s, confidence=%.2f, elapsed=%dms",
                    ocr_result.provider,
                    ocr_result.average_confidence,
                    ocr_result.total_elapsed_ms,
                )

                # Release OCR model memory (staged loading)
                if hasattr(self.ocr_engine, "unload"):
                    self.ocr_engine.unload()
            else:
                raise ParseError(
                    f"Extracted text too short ({len(text.strip())} chars). "
                    f"The file appears to be a scanned PDF. "
                    f"Enable OCR with: --ocr flag or CR_OCR_ENABLED=true. "
                    f"Install OCR dependencies: pip install -e '.[ocr]'"
                )

        if len(text.strip()) < self.MIN_TEXT_LENGTH:
            raise ParseError(
                f"Text extraction failed ({len(text.strip())} chars after OCR). "
                f"The document may be unreadable or corrupted."
            )

        sections = self._split_sections(text)

        return Contract(
            name=path.stem,
            source_path=str(path.absolute()),
            full_text=text,
            sections=sections,
            metadata=metadata,
        )

    def parse(self, file_path: str) -> Contract:
        """Parse a contract file (sync version, no OCR support).

        Raises ParseError if the file cannot be read or contains no extractable text.
        """
        path = Path(file_path)
        if not path.exists():
            raise ParseError(f"File not found: {file_path}")

        text = self._extract_text(path)
        if len(text.strip()) < self.MIN_TEXT_LENGTH:
            if self.ocr_engine:
                raise ParseError(
                    f"Scanned PDF detected. Use parse_async() for OCR support, "
                    f"or run with: contract-review {file_path} --ocr"
                )
            raise ParseError(
                f"Extracted text too short ({len(text.strip())} chars). "
                f"The file may be a scanned PDF or image-only document. "
                f"Enable OCR with: --ocr flag or CR_OCR_ENABLED=true"
            )

        sections = self._split_sections(text)

        return Contract(
            name=path.stem,
            source_path=str(path.absolute()),
            full_text=text,
            sections=sections,
        )

    def _extract_text(self, path: Path) -> str:
        """Extract raw text from file based on extension."""
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix == ".docx":
            return self._parse_docx(path)
        elif suffix in (".txt", ".md"):
            return self._read_text_file(path)
        else:
            raise ParseError(f"Unsupported file format: {suffix}")

    def _read_text_file(self, path: Path) -> str:
        """Read a text file, trying multiple encodings."""
        for encoding in ("utf-8", "gb18030", "gbk", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except (UnicodeDecodeError, ValueError):
                continue
        raise ParseError(f"Cannot decode {path.name} with any supported encoding")

    def _parse_pdf(self, path: Path) -> str:
        import pdfplumber

        try:
            with pdfplumber.open(path) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    # Also extract tables as text
                    tables = page.extract_tables() or []
                    for table in tables:
                        rows = []
                        for row in table:
                            cells = [str(c or "") for c in row]
                            rows.append(" | ".join(cells))
                        text += "\n" + "\n".join(rows)
                    pages.append(text)
                return "\n\n".join(pages)
        except Exception as e:
            raise ParseError(f"Failed to parse PDF {path.name}: {e}") from e

    def _parse_docx(self, path: Path) -> str:
        try:
            import docx

            doc = docx.Document(str(path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(cells))
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise ParseError(f"Failed to parse DOCX {path.name}: {e}") from e

    def _split_sections(self, text: str) -> list[Section]:
        """Split text into sections by detecting clause/article boundaries."""
        lines = text.split("\n")
        sections: list[Section] = []
        current_heading = ""
        current_body_lines: list[str] = []
        section_index = 0

        for line in lines:
            stripped = line.strip()
            if not stripped:
                current_body_lines.append("")
                continue

            is_heading = any(p.match(stripped) for p in self.SECTION_PATTERNS)

            if is_heading and current_body_lines:
                # Save previous section
                body = "\n".join(current_body_lines).strip()
                if body or current_heading:
                    sections.append(
                        Section(
                            heading=current_heading,
                            body=body,
                            section_type=self._detect_type(current_heading, body),
                            index=section_index,
                        )
                    )
                    section_index += 1
                current_heading = stripped
                current_body_lines = []
            elif is_heading:
                current_heading = stripped
            else:
                current_body_lines.append(stripped)

        # Save last section
        body = "\n".join(current_body_lines).strip()
        if body or current_heading:
            sections.append(
                Section(
                    heading=current_heading,
                    body=body,
                    section_type=self._detect_type(current_heading, body),
                    index=section_index,
                )
            )

        return sections if sections else [Section(body=text, section_type="body")]

    @staticmethod
    def _detect_type(heading: str, body: str) -> str:
        """Heuristic section type detection."""
        combined = (heading + " " + body[:200]).lower()
        if any(kw in combined for kw in ["甲方", "乙方", "party a", "party b", "签署方"]):
            return "parties"
        if any(kw in combined for kw in ["签章", "签字", "盖章", "signature", "signed by"]):
            return "signature"
        if "|" in body and body.count("|") > 3:
            return "table"
        return "body"
